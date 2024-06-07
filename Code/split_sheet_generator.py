import os
import pandas as pd
import docx
from fuzzywuzzy import process
from datetime import datetime

def load_config(config_path):
    try:
        config = pd.read_excel(config_path)
        config_dict = {row['Key']: row['Value'] for _, row in config.iterrows()}
        return config_dict
    except KeyError as e:
        print(f"Error: Missing column in the configuration file - {e}")
        raise
    except Exception as e:
        print(f"An error occurred while loading the configuration file: {e}")
        raise

def get_user_input(prompt, optional=False):
    user_input = input(prompt)
    if optional and user_input.strip() == "":
        return None
    return user_input.strip()

def copy_styles(source_doc, target_doc):
    styles = source_doc.styles
    for style in styles:
        if style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH:
            if style.name not in target_doc.styles:
                new_style = target_doc.styles.add_style(style.name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                new_style.base_style = style.base_style
                new_style.font.name = style.font.name
                new_style.font.size = style.font.size
                new_style.font.bold = style.font.bold
                new_style.font.italic = style.font.italic
                new_style.font.underline = style.font.underline
                new_style.paragraph_format.alignment = style.paragraph_format.alignment

def create_new_document(template_path, text_path, add_date, output_dir, company_name, song_name):
    try:
        # Load template and text documents
        template_doc = docx.Document(template_path)
        text_doc = docx.Document(text_path)
    except Exception as e:
        print(f"Error: {e}")
        return None, None

    # Create new document and copy styles
    new_doc = docx.Document()
    copy_styles(template_doc, new_doc)

    # Add company name
    new_doc.add_paragraph(company_name, style='Title')

    # Add date if required
    if add_date:
        current_date = datetime.now()
        day = current_date.day
        month = current_date.strftime('%B')
        year = current_date.year
        date_paragraph = f"Dated the {day} day of {month}, {year}"
        new_doc.add_paragraph(date_paragraph)

    # Add text from text_doc
    for para in text_doc.paragraphs:
        new_doc.add_paragraph(para.text, style=para.style.name)

    # Save the new document with a unique name
    version = 1
    base_name = song_name if song_name else "Untitled"
    new_filename = f"{base_name}_V{version}_Split_Sheet.docx"
    new_filepath = os.path.join(output_dir, new_filename)
    while os.path.exists(new_filepath):
        version += 1
        new_filename = f"{base_name}_V{version}_Split_Sheet.docx"
        new_filepath = os.path.join(output_dir, new_filename)
    new_doc.save(new_filepath)
    return new_filepath, new_doc

def update_document_with_song_info(doc, song_name, artist_name):
    if song_name:
        for paragraph in doc.paragraphs:
            if "SONG TITLE" in paragraph.text:
                paragraph.text = paragraph.text.replace("SONG TITLE", f'"{song_name}"')
    if artist_name:
        for paragraph in doc.paragraphs:
            if "ARTIST NAME" in paragraph.text:
                paragraph.text = paragraph.text.replace("ARTIST NAME", f'"{artist_name}"')

def load_excel_data(excel_path):
    return pd.read_excel(excel_path)

def find_matching_writers(data, writer_name):
    matches = []
    writer_name_lower = writer_name.lower()
    for index, row in data.iterrows():
        name = row['Writer Name']
        ratio = process.extractOne(writer_name_lower, [name.lower()])
        if ratio and ratio[1] >= 75:
            matches.append((index, row))
    return matches

def display_matches(matches):
    for i, (index, row) in enumerate(matches):
        print(f"{i + 1}: {row.to_dict()}")

def insert_table_with_signatures(doc, writer_list):
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'WRITERS'
    hdr_cells[1].text = 'PUBLISHERS'
    hdr_cells[2].text = 'OWNERSHIP (%)'
    hdr_cells[3].text = 'PRO'
    hdr_cells[4].text = 'SIGNATURE'

    for writer in writer_list:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{writer['Writer Name']}\nIPI: {writer['Writer IPI']}"
        row_cells[1].text = f"{writer['Publisher Name']}\nIPI: {writer['Publisher IPI']}"
        row_cells[2].text = ""
        row_cells[3].text = writer['Writer PRO']
        row_cells[4].text = ""

    # Add spacing between the table and the signature boxes
    doc.add_paragraph("\n")

    # Insert signature boxes with horizontal alignment
    signature_box_per_row = 3
    for i in range(0, len(writer_list), signature_box_per_row):
        row = doc.add_paragraph()
        for j in range(signature_box_per_row):
            if i + j < len(writer_list):
                writer = writer_list[i + j]
                signature_box = f"_____________________________\nSignature: {writer['Writer Name']}"
                run = row.add_run(signature_box)
                run.add_break(docx.enum.text.WD_BREAK.LINE)
                run.add_break(docx.enum.text.WD_BREAK.LINE)  # Add more space between signature boxes vertically

def main():
    config_path = r"C:\Users\alain\Alain Developer\General Projects\Split Sheet Generator\Documentation\Config_SplitSheetGenerator.xlsx"
    config = load_config(config_path)
    
    company_name = config.get("Company Name", "Company Name Not Specified")
    template_path = config.get("Template Path")
    text_path = config.get("Text Path")
    excel_path = config.get("Excel Path")
    output_dir = config.get("Output Directory")
    
    add_date_choice = get_user_input("Press 1 to add date or 0 to continue: ")
    add_date = add_date_choice == "1"
    
    song_name = get_user_input("Enter song name: ", optional=True)
    if not song_name:
        print("Please remember to specify a song title in the split sheet at a later date")
    artist_name = get_user_input("Please enter artist name: ", optional=True)
    
    new_filepath, new_doc = create_new_document(template_path, text_path, add_date, output_dir, company_name, song_name)
    if new_filepath is None or new_doc is None:
        print("Error creating the document. Please check the file paths and try again.")
        return

    update_document_with_song_info(new_doc, song_name, artist_name)
    new_doc.save(new_filepath)
    
    writer_data = load_excel_data(excel_path)
    writer_list = []

    while True:
        writer_name = get_user_input("Enter writer's name or press 0 to finish: ")
        if writer_name == "0":
            insert_table_with_signatures(new_doc, writer_list)
            new_doc.save(new_filepath)
            print(f"Thank you, you may find your template here: {new_filepath}")
            break
        
        if any(writer['Writer Name'].lower() == writer_name.lower() for writer in writer_list):
            print("This writer has already been added, please choose another.")
            continue

        matches = find_matching_writers(writer_data, writer_name)
        if matches:
            display_matches(matches)
            choice = get_user_input("Enter the number next to the writer you want to add, press 9 to go back to writer search, or press 0 to exit: ")
            if choice == "0":
                insert_table_with_signatures(new_doc, writer_list)
                new_doc.save(new_filepath)
                print(f"Thank you, you may find your template here: {new_filepath}")
                break
            elif choice == "9":
                continue
            else:
                try:
                    selected_index = int(choice) - 1
                    if 0 <= selected_index < len(matches):
                        selected_writer = matches[selected_index][1]
                        writer_list.append({
                            'Writer Name': selected_writer['Writer Name'],
                            'Writer IPI': selected_writer['Writer IPI'],
                            'Publisher Name': selected_writer['Publisher Name'],
                            'Publisher IPI': selected_writer['Publisher IPI'],
                            'Writer PRO': selected_writer['Writer PRO']
                        })
                    else:
                        print("Invalid selection. Please try again.")
                except ValueError:
                    print("Invalid input. Please enter a number.")
        else:
            print("No matches found. Please try again.")

if __name__ == "__main__":
    main()
